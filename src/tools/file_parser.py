"""
File Parser for Trip Planning References

Supports parsing various file formats to extract trip planning information:
- Text files (.txt)
- PDF documents (.pdf)
- Excel spreadsheets (.xlsx)
- Word documents (.docx)
- Google Maps links
- Wanderlog links

This allows users to upload friends' trip plans and use them as reference!
"""

from pathlib import Path
from typing import Dict, List, Optional
import re


class FileParser:
    """
    Parse reference files for trip planning context.
    
    Supports:
    - .txt files (plain text itineraries)
    - .pdf files (travel documents, printouts)
    - .xlsx files (budget spreadsheets, day-by-day plans)
    - .docx files (detailed trip plans)
    - Google Maps links (saved locations)
    - Wanderlog links (shared itineraries)
    """
    
    def __init__(self):
        self.supported_formats = ['.txt', '.pdf', '.xlsx', '.docx']
    
    def parse_file(self, file_path: str) -> Dict[str, any]:
        """
        Parse a reference file and extract trip planning information.
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            Dictionary with extracted information:
            {
                'file_name': str,
                'file_type': str,
                'content': str,
                'structured_data': dict (if available),
                'links': list (if found),
                'destinations': list (if identified),
                'dates': list (if found),
                'costs': list (if found)
            }
        """
        path = Path(file_path)
        
        if not path.exists():
            return {
                'error': f"File not found: {file_path}",
                'file_name': path.name,
                'content': ''
            }
        
        file_type = path.suffix.lower()
        
        if file_type == '.txt':
            return self._parse_txt(path)
        elif file_type == '.pdf':
            return self._parse_pdf(path)
        elif file_type == '.xlsx':
            return self._parse_xlsx(path)
        elif file_type == '.docx':
            return self._parse_docx(path)
        else:
            return {
                'error': f"Unsupported file type: {file_type}",
                'file_name': path.name,
                'supported': self.supported_formats
            }
    
    def _parse_txt(self, path: Path) -> Dict:
        """Parse plain text file"""
        try:
            content = path.read_text(encoding='utf-8')
            
            # Check if it's a link file
            if self._is_google_maps_link(content):
                return self._parse_google_maps_link(content, path.name)
            elif self._is_wanderlog_link(content):
                return self._parse_wanderlog_link(content, path.name)
            
            # Regular text file
            return {
                'file_name': path.name,
                'file_type': 'text',
                'content': content,
                'links': self._extract_links(content),
                'destinations': self._extract_destinations(content),
                'dates': self._extract_dates(content),
                'costs': self._extract_costs(content),
                'activities': self._extract_activities(content)
            }
        except Exception as e:
            return {
                'error': f"Error parsing text file: {e}",
                'file_name': path.name,
                'content': ''
            }
    
    def _parse_pdf(self, path: Path) -> Dict:
        """Parse PDF file"""
        try:
            # Try to import PDF parser
            try:
                import PyPDF2
                with open(path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    content = ""
                    for page in reader.pages:
                        content += page.extract_text() + "\n"
            except ImportError:
                content = f"[PDF: {path.name}] - PyPDF2 not installed. Install with: pip install PyPDF2"
            
            return {
                'file_name': path.name,
                'file_type': 'pdf',
                'content': content,
                'links': self._extract_links(content),
                'destinations': self._extract_destinations(content),
                'dates': self._extract_dates(content),
                'costs': self._extract_costs(content)
            }
        except Exception as e:
            return {
                'error': f"Error parsing PDF: {e}",
                'file_name': path.name,
                'content': f"[PDF: {path.name}] - Could not parse"
            }
    
    def _parse_xlsx(self, path: Path) -> Dict:
        """Parse Excel file"""
        try:
            # Try to import Excel parser
            try:
                import openpyxl
                workbook = openpyxl.load_workbook(path)
                
                content = ""
                structured_data = {}
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    content += f"\n=== {sheet_name} ===\n"
                    
                    sheet_data = []
                    for row in sheet.iter_rows(values_only=True):
                        if any(cell is not None for cell in row):
                            row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                            content += row_str + "\n"
                            sheet_data.append(row)
                    
                    structured_data[sheet_name] = sheet_data
                
            except ImportError:
                content = f"[Excel: {path.name}] - openpyxl not installed. Install with: pip install openpyxl"
                structured_data = {}
            
            return {
                'file_name': path.name,
                'file_type': 'excel',
                'content': content,
                'structured_data': structured_data,
                'costs': self._extract_costs(content)
            }
        except Exception as e:
            return {
                'error': f"Error parsing Excel: {e}",
                'file_name': path.name,
                'content': f"[Excel: {path.name}] - Could not parse"
            }
    
    def _parse_docx(self, path: Path) -> Dict:
        """Parse Word document"""
        try:
            # Try to import Word parser
            try:
                import docx
                doc = docx.Document(path)
                
                content = ""
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                
                # Also extract tables
                for table in doc.tables:
                    content += "\n[Table]\n"
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        content += row_text + "\n"
            except ImportError:
                content = f"[Word: {path.name}] - python-docx not installed. Install with: pip install python-docx"
            
            return {
                'file_name': path.name,
                'file_type': 'word',
                'content': content,
                'links': self._extract_links(content),
                'destinations': self._extract_destinations(content),
                'dates': self._extract_dates(content),
                'costs': self._extract_costs(content)
            }
        except Exception as e:
            return {
                'error': f"Error parsing Word document: {e}",
                'file_name': path.name,
                'content': f"[Word: {path.name}] - Could not parse"
            }
    
    def _is_google_maps_link(self, content: str) -> bool:
        """Check if content is a Google Maps link"""
        return 'google.com/maps' in content.lower() or 'goo.gl/maps' in content.lower()
    
    def _is_wanderlog_link(self, content: str) -> bool:
        """Check if content is a Wanderlog link"""
        return 'wanderlog.com' in content.lower()
    
    def _parse_google_maps_link(self, content: str, filename: str) -> Dict:
        """Parse Google Maps link"""
        links = re.findall(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', content)
        
        return {
            'file_name': filename,
            'file_type': 'google_maps_link',
            'content': content,
            'links': links,
            'note': 'Google Maps link detected - agent can use this to find locations and attractions'
        }
    
    def _parse_wanderlog_link(self, content: str, filename: str) -> Dict:
        """Parse Wanderlog link"""
        links = re.findall(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', content)
        
        return {
            'file_name': filename,
            'file_type': 'wanderlog_link',
            'content': content,
            'links': links,
            'note': 'Wanderlog link detected - agent can reference this shared itinerary'
        }
    
    def _extract_links(self, content: str) -> List[str]:
        """Extract URLs from content"""
        pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        return re.findall(pattern, content)
    
    def _extract_destinations(self, content: str) -> List[str]:
        """Extract destination names (simple pattern matching)"""
        # Look for patterns like "Visit X", "Trip to X", city names in caps
        destinations = []
        patterns = [
            r'(?:visit|trip to|going to|destination:)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s+(?:Japan|France|Italy|UK|USA|Spain|Germany|China)',
        ]
        for pattern in patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            destinations.extend(matches)
        return list(set(destinations))[:10]  # Limit to top 10 unique
    
    def _extract_dates(self, content: str) -> List[str]:
        """Extract dates (various formats)"""
        patterns = [
            r'\d{4}-\d{2}-\d{2}',  # YYYY-MM-DD
            r'\d{2}/\d{2}/\d{4}',  # DD/MM/YYYY or MM/DD/YYYY
            r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}',  # Month DD, YYYY
        ]
        dates = []
        for pattern in patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            dates.extend(matches)
        return list(set(dates))[:20]  # Limit to 20 dates
    
    def _extract_costs(self, content: str) -> List[str]:
        """Extract cost/price information"""
        patterns = [
            r'\$\d+(?:,\d{3})*(?:\.\d{2})?',  # $1,234.56
            r'€\d+(?:,\d{3})*(?:\.\d{2})?',  # €1,234.56
            r'¥\d+(?:,\d{3})*',  # ¥1,234
            r'£\d+(?:,\d{3})*(?:\.\d{2})?',  # £1,234.56
        ]
        costs = []
        for pattern in patterns:
            matches = re.findall(pattern, content)
            costs.extend(matches)
        return list(set(costs))[:50]  # Limit to 50 costs
    
    def _extract_activities(self, content: str) -> List[str]:
        """Extract activity mentions"""
        # Look for common activity keywords
        activity_keywords = [
            'visit', 'see', 'tour', 'explore', 'hike', 'museum', 'temple', 
            'shrine', 'restaurant', 'cafe', 'shopping', 'market', 'park',
            'beach', 'mountain', 'castle', 'garden'
        ]
        activities = []
        for keyword in activity_keywords:
            pattern = rf'{keyword}\s+([a-zA-Z\s]+?)(?:\.|,|\n|$)'
            matches = re.findall(pattern, content, re.IGNORECASE)
            activities.extend([f"{keyword.title()} {match.strip()}" for match in matches if len(match.strip()) > 3])
        return list(set(activities))[:30]  # Limit to 30 activities


def parse_reference_files(file_paths: List[str]) -> List[Dict]:
    """
    Parse multiple reference files.
    
    Args:
        file_paths: List of file paths to parse
        
    Returns:
        List of parsed file data dictionaries
    """
    parser = FileParser()
    results = []
    
    for file_path in file_paths:
        result = parser.parse_file(file_path)
        results.append(result)
    
    return results


def create_reference_context(parsed_files: List[Dict]) -> str:
    """
    Create a context string from parsed reference files for the agent.
    
    Args:
        parsed_files: List of parsed file dictionaries
        
    Returns:
        Formatted context string for the agent
    """
    if not parsed_files:
        return ""
    
    context = "=== REFERENCE FILES FROM FRIENDS' TRIPS ===\n\n"
    
    for i, file_data in enumerate(parsed_files, 1):
        context += f"Reference {i}: {file_data.get('file_name', 'Unknown')}\n"
        context += f"Type: {file_data.get('file_type', 'unknown')}\n"
        
        if 'error' in file_data:
            context += f"Note: {file_data['error']}\n\n"
            continue
        
        # Add content summary
        content = file_data.get('content', '')
        if content:
            # Truncate if too long
            max_length = 2000
            if len(content) > max_length:
                content = content[:max_length] + f"\n... (truncated, {len(content)} total chars)"
            context += f"Content:\n{content}\n"
        
        # Add extracted information
        if file_data.get('destinations'):
            context += f"Destinations mentioned: {', '.join(file_data['destinations'])}\n"
        
        if file_data.get('dates'):
            context += f"Dates found: {', '.join(file_data['dates'][:5])}\n"
        
        if file_data.get('costs'):
            context += f"Costs mentioned: {', '.join(file_data['costs'][:10])}\n"
        
        if file_data.get('activities'):
            context += f"Activities: {', '.join(file_data['activities'][:10])}\n"
        
        if file_data.get('links'):
            context += f"Links: {', '.join(file_data['links'][:5])}\n"
        
        if file_data.get('note'):
            context += f"Note: {file_data['note']}\n"
        
        context += "\n" + "─" * 60 + "\n\n"
    
    context += "=== END OF REFERENCE FILES ===\n"
    return context

